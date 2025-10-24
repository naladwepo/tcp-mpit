"""
Генератор технико-коммерческих предложений (ТКП) в форматах Word и PDF
"""

from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import locale

try:
    locale.setlocale(locale.LC_TIME, 'ru_RU.UTF-8')
except:
    try:
        locale.setlocale(locale.LC_TIME, 'ru_RU')
    except:
        pass  # Используем системную локаль


class DocumentGenerator:
    """Генератор документов ТКП"""
    
    # Данные компании
    COMPANY_INFO = {
        "name": "ООО «СтройТехКомплект»",
        "address": "123456, г. Москва, ул. Промышленная, д. 10, офис 305",
        "phone": "+7 (495) 123-45-67",
        "email": "info@stroytechcomplex.ru",
        "inn": "7701234567",
        "kpp": "770101001",
        "ogrn": "1127746123456",
        "director": "Иванов Иван Иванович"
    }
    
    def __init__(self, output_dir: str = "generated_documents"):
        """
        Инициализация генератора
        
        Args:
            output_dir: Директория для сохранения документов
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True, parents=True)
    
    def _add_border_to_cell(self, cell):
        """Добавляет границы к ячейке таблицы"""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        
        # Создаем границы
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        
        tcPr.append(tcBorders)
    
    def _format_currency(self, amount: float) -> str:
        """Форматирует сумму в рублях"""
        return f"{amount:,.2f}".replace(",", " ").replace(".", ",") + " ₽"
    
    def _add_company_header(self, doc: Document):
        """Добавляет шапку с данными компании"""
        # Заголовок
        title = doc.add_heading('ТЕХНИКО-КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Номер и дата
        date_str = datetime.now().strftime("%d.%m.%Y")
        doc_number = f"№ ТКП-{datetime.now().strftime('%Y%m%d-%H%M')}"
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"{doc_number}\nот {date_str}")
        run.font.size = Pt(11)
        
        doc.add_paragraph()  # Пустая строка
        
        # Данные компании
        company_table = doc.add_table(rows=7, cols=2)
        company_table.style = 'Light List Accent 1'
        
        company_data = [
            ("Поставщик:", self.COMPANY_INFO["name"]),
            ("Адрес:", self.COMPANY_INFO["address"]),
            ("Телефон:", self.COMPANY_INFO["phone"]),
            ("Email:", self.COMPANY_INFO["email"]),
            ("ИНН/КПП:", f"{self.COMPANY_INFO['inn']} / {self.COMPANY_INFO['kpp']}"),
            ("ОГРН:", self.COMPANY_INFO["ogrn"]),
            ("Директор:", self.COMPANY_INFO["director"])
        ]
        
        for i, (label, value) in enumerate(company_data):
            row = company_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Форматирование
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                if cell == row.cells[0]:
                    cell.paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()  # Пустая строка
    
    def _add_items_table(self, doc: Document, items: List[Dict[str, Any]], query: str):
        """Добавляет таблицу с товарами"""
        # Заголовок таблицы
        p = doc.add_paragraph()
        run = p.add_run(f"Запрос: {query}")
        run.font.bold = True
        run.font.size = Pt(11)
        
        doc.add_paragraph()  # Пустая строка
        
        # Создаем таблицу
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        # Заголовки
        headers = ['№', 'Наименование товара', 'Спецификация', 'Кол-во', 'Цена за ед.', 'Сумма']
        header_cells = table.rows[0].cells
        
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = header_text
            
            # Форматирование заголовка
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0]
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Цвет фона заголовка
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '4472C4')
            cell._element.get_or_add_tcPr().append(shading)
            
            self._add_border_to_cell(cell)
        
        # Заполняем данные
        total_cost = 0
        row_num = 1
        
        for item in items:
            # Пропускаем товары, которые не найдены
            if not item.get('found_product'):
                continue
            
            product = item['found_product']
            row = table.add_row()
            
            # Данные строки
            row.cells[0].text = str(row_num)
            row.cells[1].text = product.get('name', 'Не указано')
            row.cells[2].text = item.get('specifications', '-')
            row.cells[3].text = str(item.get('quantity', 1))
            row.cells[4].text = self._format_currency(item.get('unit_price', 0))
            row.cells[5].text = self._format_currency(item.get('total_price', 0))
            
            total_cost += item.get('total_price', 0)
            
            # Форматирование ячеек
            for j, cell in enumerate(row.cells):
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.font.size = Pt(9)
                
                # Выравнивание
                if j in [0, 3]:  # Номер и количество - по центру
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif j in [4, 5]:  # Цены - по правому краю
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                self._add_border_to_cell(cell)
            
            row_num += 1
        
        # Итоговая строка
        row = table.add_row()
        row.cells[0].merge(row.cells[4])
        row.cells[0].text = "ИТОГО:"
        row.cells[1].text = self._format_currency(total_cost)
        
        # Форматирование итоговой строки
        for cell in [row.cells[0], row.cells[1]]:
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0]
            run.font.bold = True
            run.font.size = Pt(11)
            
            if cell == row.cells[0]:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Цвет фона
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'E7E6E6')
            cell._element.get_or_add_tcPr().append(shading)
            
            self._add_border_to_cell(cell)
        
        return total_cost
    
    def _add_footer(self, doc: Document, total_cost: float):
        """Добавляет подвал документа"""
        doc.add_paragraph()  # Пустая строка
        
        # Итоговая сумма прописью
        p = doc.add_paragraph()
        run = p.add_run(f"Общая сумма: {self._format_currency(total_cost)}")
        run.font.bold = True
        run.font.size = Pt(12)
        
        doc.add_paragraph()  # Пустая строка
        
        # Условия
        p = doc.add_paragraph()
        run = p.add_run("Условия поставки:")
        run.font.bold = True
        run.font.size = Pt(11)
        
        conditions = [
            "Срок поставки: 10-14 рабочих дней с момента оплаты",
            "Условия оплаты: 100% предоплата по безналичному расчету",
            "Гарантия: 12 месяцев от производителя",
            "Доставка: по согласованию (стоимость рассчитывается отдельно)",
            "Цены действительны в течение 14 календарных дней"
        ]
        
        for condition in conditions:
            p = doc.add_paragraph(condition, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            for run in p.runs:
                run.font.size = Pt(10)
        
        doc.add_paragraph()  # Пустая строка
        
        # Подпись
        doc.add_paragraph()
        doc.add_paragraph()
        
        signature_table = doc.add_table(rows=1, cols=2)
        signature_table.rows[0].cells[0].text = f"Директор\n{self.COMPANY_INFO['director']}"
        signature_table.rows[0].cells[1].text = "_____________ / ___________"
        
        for cell in signature_table.rows[0].cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
        
        signature_table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def generate_word(self, data: Dict[str, Any], filename: str = None) -> Path:
        """
        Генерирует документ Word
        
        Args:
            data: Данные для документа (результат поиска)
            filename: Имя файла (опционально)
            
        Returns:
            Path: Путь к созданному файлу
        """
        # Создаем документ
        doc = Document()
        
        # Настройки документа
        sections = doc.sections
        for section in sections:
            section.page_height = Inches(11.69)  # A4
            section.page_width = Inches(8.27)
            section.left_margin = Inches(0.79)
            section.right_margin = Inches(0.59)
            section.top_margin = Inches(0.79)
            section.bottom_margin = Inches(0.79)
        
        # Шапка компании
        self._add_company_header(doc)
        
        # Таблица с товарами
        total_cost = self._add_items_table(
            doc, 
            data.get('items', []),
            data.get('original_query', 'Не указан')
        )
        
        # Подвал
        self._add_footer(doc, total_cost)
        
        # Сохраняем
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"TKP_{timestamp}.docx"
        
        filepath = self.output_dir / filename
        doc.save(str(filepath))
        
        return filepath
    
    def generate_pdf(self, data: Dict[str, Any], filename: str = None) -> Path:
        """
        Генерирует документ PDF через LibreOffice или Word
        
        Args:
            data: Данные для документа (результат поиска)
            filename: Имя файла (опционально)
            
        Returns:
            Path: Путь к созданному файлу
        """
        # Сначала создаем Word документ
        if filename:
            word_filename = filename.replace('.pdf', '.docx')
        else:
            word_filename = None
        
        word_path = self.generate_word(data, word_filename)
        
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"TKP_{timestamp}.pdf"
        
        pdf_path = self.output_dir / filename
        
        # Пробуем конвертировать через LibreOffice (предпочтительно на macOS/Linux)
        import subprocess
        import sys
        import shutil
        
        # Проверяем наличие LibreOffice
        libreoffice_paths = [
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS
            '/usr/bin/soffice',  # Linux
            '/usr/bin/libreoffice',  # Linux альтернатива
            shutil.which('soffice'),  # Поиск в PATH
            shutil.which('libreoffice'),  # Поиск в PATH
        ]
        
        libreoffice_cmd = None
        for path in libreoffice_paths:
            if path and Path(path).exists():
                libreoffice_cmd = path
                break
        
        if libreoffice_cmd:
            try:
                # Конвертируем через LibreOffice (без запросов доступа!)
                subprocess.run([
                    libreoffice_cmd,
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', str(self.output_dir),
                    str(word_path)
                ], check=True, capture_output=True)
                
                return pdf_path
                
            except subprocess.CalledProcessError as e:
                print(f"⚠️  Ошибка LibreOffice: {e}")
                print("   Пробуем альтернативный метод...")
        
        # Если LibreOffice не доступен, пробуем docx2pdf (использует Word на macOS)
        try:
            from docx2pdf import convert
            convert(str(word_path), str(pdf_path))
            return pdf_path
            
        except ImportError:
            print("⚠️  Библиотека docx2pdf не установлена и LibreOffice не найден.")
            print("   Установите LibreOffice для конвертации без запросов доступа:")
            print("   brew install --cask libreoffice")
            print("   Или установите: pip install docx2pdf (использует Word)")
            return word_path
        except Exception as e:
            print(f"⚠️  Ошибка при создании PDF: {e}")
            print("   Создан только Word документ")
            return word_path
    
    def generate_both(self, data: Dict[str, Any], base_filename: str = None) -> Dict[str, Path]:
        """
        Генерирует оба формата (Word и PDF)
        
        Args:
            data: Данные для документа (результат поиска)
            base_filename: Базовое имя файла без расширения
            
        Returns:
            Dict[str, Path]: Словарь с путями к созданным файлам
        """
        if base_filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_filename = f"TKP_{timestamp}"
        
        word_filename = f"{base_filename}.docx"
        pdf_filename = f"{base_filename}.pdf"
        
        word_path = self.generate_word(data, word_filename)
        pdf_path = self.generate_pdf(data, pdf_filename)
        
        return {
            'word': word_path,
            'pdf': pdf_path
        }


# Пример использования
if __name__ == "__main__":
    # Тестовые данные
    test_data = {
        "original_query": "Комплект для монтажа короба 200x200",
        "items": [
            {
                "requested_item": "Короб 200x200",
                "quantity": 2,
                "found_product": {
                    "name": "Короб 200x200 мм, L=2000 мм, горячее цинкование",
                    "cost": 88498.0
                },
                "specifications": "200x200 мм, L=2000 мм",
                "unit_price": 88498.0,
                "total_price": 176996.0
            },
            {
                "requested_item": "Крышка для короба",
                "quantity": 2,
                "found_product": {
                    "name": "Крышка 200 мм, L=2000 мм, горячее цинкование",
                    "cost": 45131.0
                },
                "specifications": "200 мм, L=2000 мм",
                "unit_price": 45131.0,
                "total_price": 90262.0
            },
            {
                "requested_item": "Винты М6",
                "quantity": 10,
                "found_product": {
                    "name": "Винт с крестообразным шлицем М6х10",
                    "cost": 100.0
                },
                "specifications": "М6х10",
                "unit_price": 100.0,
                "total_price": 1000.0
            }
        ]
    }
    
    generator = DocumentGenerator()
    files = generator.generate_both(test_data)
    
    print("✅ Документы созданы:")
    print(f"   Word: {files['word']}")
    print(f"   PDF: {files['pdf']}")
